import pandas as pd
import numpy as np
import io
import faiss
import pickle
import os
import asyncio
from typing import List, Dict, Any
import openai
from openai import AsyncOpenAI
import streamlit as st
import nest_asyncio

# Enable nested asyncio for environments like Jupyter/Spyder
nest_asyncio.apply()

# Define RAG system class to store embeddings, vector database, and other components
class RAGSystem:
    def __init__(self):
        self.documents = []
        self.embeddings = []
        self.document_metadata = {}
        self.faiss_index = None
        self.page_metadata = []
        self.embedding_model = "text-embedding-3-large"
        self.is_initialized = False
        self.index_path = "faiss_index"

def load_faiss_index(index_path: str = "faiss_index"):
    """
    Load existing FAISS index and metadata.
    
    Args:
        index_path (str): Path to the saved index directory
        
    Returns:
        tuple: (faiss_index, page_metadata)
    """
    try:
        # Load FAISS index
        index = faiss.read_index(os.path.join(index_path, "faiss.index"))
        
        # Load metadata
        with open(os.path.join(index_path, "metadata.pkl"), 'rb') as f:
            page_metadata = pickle.load(f)
        
        
        return index, page_metadata
        
    except FileNotFoundError as e:
        print(f"Error: Could not load index from '{index_path}': {e}")
        print("Make sure to run generate_embeddings.py first to build the index")
        return None, None

async def get_query_embedding_async(openai_api_key: str, query: str, embedding_model: str) -> np.ndarray:
    """
    Get embedding for a query using async OpenAI API.
    
    Args:
        openai_api_key (str): OpenAI API key
        query (str): Query text to embed
        embedding_model (str): OpenAI embedding model to use
        
    Returns:
        np.ndarray: Query embedding vector
    """
    client = AsyncOpenAI(api_key=openai_api_key)
    
    try:
        response = await client.embeddings.create(
            input=query,
            model=embedding_model
        )
        embedding = np.array(response.data[0].embedding, dtype=np.float32)
        # Normalize for cosine similarity
        embedding = embedding / np.linalg.norm(embedding)
        return embedding
    finally:
        await client.close()

def get_query_embedding(openai_api_key: str, query: str, embedding_model: str) -> np.ndarray:
    """
    Synchronous wrapper for getting query embedding.
    
    Args:
        openai_api_key (str): OpenAI API key
        query (str): Query text to embed
        embedding_model (str): OpenAI embedding model to use
        
    Returns:
        np.ndarray: Query embedding vector
    """
    try:
        return asyncio.run(get_query_embedding_async(openai_api_key, query, embedding_model))
    except RuntimeError as e:
        if "cannot be called from a running event loop" in str(e):
            loop = asyncio.get_event_loop()
            return loop.run_until_complete(get_query_embedding_async(openai_api_key, query, embedding_model))
        else:
            raise e

def initialize_rag_system(api_key: str, document_data: pd.DataFrame = None, index_path: str = "faiss_index") -> RAGSystem:
    """
    Initialize the RAG system with FAISS index
    
    Parameters:
    -----------
    api_key : str
        OpenAI API key for generating embeddings
    document_data : pd.DataFrame, optional
        Dataframe containing document data (not used in FAISS implementation)
    index_path : str
        Path to the FAISS index directory
    
    Returns:
    --------
    RAGSystem
        Initialized RAG system
    """
    # Set OpenAI API key
    openai.api_key = api_key
    
    # Create a new RAG system
    rag_system = RAGSystem()
    rag_system.index_path = index_path
    
    # Load FAISS index and metadata
    index, page_metadata = load_faiss_index(index_path)
    
    if index is None or page_metadata is None:
        st.error(f"Failed to load FAISS index from '{index_path}'. Please ensure the index exists.")
        return rag_system
    
    # Store in RAG system
    rag_system.faiss_index = index
    rag_system.page_metadata = page_metadata
    rag_system.is_initialized = True
    
    # Convert page metadata to documents format for compatibility
    rag_system.documents = []
    for i, page in enumerate(page_metadata):
        doc = {
            'id': i,
            'title': f"{page['pdf_name']} - Page {page['page_number']}",
            'text': page['content'],
            'metadata': {
                'pdf_name': page['pdf_name'],
                'page_number': page['page_number'],
                'publication_year': page['metadata'].get('Publication Year'),
                'publication_type': page['metadata'].get('Publication Type'),
                'authors': page['metadata'].get('Authors'),
                'journal': page['metadata'].get('Journal'),
                'doi': page['metadata'].get('DOI'),
                'source': page['pdf_name']
            }
        }
        rag_system.documents.append(doc)
    
    st.success(f"RAG system initialized successfully with {len(page_metadata)} documents!")
    
    return rag_system

def search_documents_faiss(query: str, rag_system: RAGSystem, openai_api_key: str, top_k: int = 8) -> List[Dict]:
    """
    Search for similar pages using the FAISS index.
    
    Args:
        query (str): Search query
        rag_system (RAGSystem): Initialized RAG system
        openai_api_key (str): OpenAI API key
        top_k (int): Number of top results to return
        
    Returns:
        List[Dict]: List of search results with metadata and scores
    """
    if not rag_system.is_initialized or rag_system.faiss_index is None:
        print("Error: RAG system not properly initialized.")
        return []
    
    
    try:
        # Get query embedding
        query_embedding = get_query_embedding(openai_api_key, query, rag_system.embedding_model)
        query_embedding = query_embedding.reshape(1, -1)
        
        # Search in FAISS index
        scores, indices = rag_system.faiss_index.search(query_embedding, top_k)
        
        # Prepare results
        results = []
        for i, (score, idx) in enumerate(zip(scores[0], indices[0])):
            if idx < len(rag_system.page_metadata):
                page_data = rag_system.page_metadata[idx]
                result = {
                    "rank": i + 1,
                    "score": float(score),
                    "pdf_name": page_data["pdf_name"],
                    "page_number": page_data["page_number"],
                    "content": page_data["content"][:500] + "..." if len(page_data["content"]) > 500 else page_data["content"],
                    "full_content": page_data["content"],
                    "metadata": page_data["metadata"]
                }
                results.append(result)
        
        return results
    
    except Exception as e:
        print(f"Error during search: {e}")
        return []

def get_relevant_documents(question: str, rag_system: RAGSystem, top_k: int = 8) -> List[Dict[str, Any]]:
    """
    Retrieve relevant documents for a given question using FAISS
    
    Parameters:
    -----------
    question : str
        User question
    rag_system : RAGSystem
        Initialized RAG system
    top_k : int
        Number of top documents to retrieve
    
    Returns:
    --------
    List[Dict[str, Any]]
        List of relevant documents with metadata
    """
    if not rag_system.is_initialized:
        return []
    
    # Get API key from session state
    api_key = st.session_state.get('openai_api_key', '')
    if not api_key:
        st.error("OpenAI API key not found in session state")
        return []
    
    # Search using FAISS
    search_results = search_documents_faiss(question, rag_system, api_key, top_k)
    
    # Format results for display
    formatted_results = []
    for result in search_results:
        metadata = result.get('metadata', {})
        
        # Create excerpt from content
        content = result.get('full_content', '')
        excerpt = content[:300] + "..." if len(content) > 300 else content
        
        formatted_result = {
            'title': f"{result['pdf_name']} - Page {result['page_number']}",
            'excerpt': excerpt,
            'score': round(result['score'], 3),
            'year': metadata.get('Publication Year', 'N/A'),
            'source': result['pdf_name'],
            'page_number': result['page_number'],
            'authors': metadata.get('Authors', 'N/A'),
            'journal': metadata.get('Journal', 'N/A'),
            'doi': metadata.get('DOI', 'N/A'),
            'full_content': result['full_content']
        }
        formatted_results.append(formatted_result)
    
    return formatted_results

async def generate_answer_async(question: str, relevant_documents: List[Dict[str, Any]], api_key: str) -> str:
    """
    Generate an answer using OpenAI API with relevant document context.
    
    Parameters:
    -----------
    question : str
        User question
    relevant_documents : List[Dict[str, Any]]
        List of relevant documents
    api_key : str
        OpenAI API key
    
    Returns:
    --------
    str
        Generated answer
    """
    if not relevant_documents:
        return "I couldn't find any relevant information in the research documents. Please try a different question."
    
    # Prepare context from relevant documents
    context_parts = []
    for i, doc in enumerate(relevant_documents):  # Use top 3 documents
        context_parts.append(f"""
Document {i+1}: {doc['title']}
Source: {doc['source']} (Page {doc['page_number']})
Authors: {doc.get('authors', 'N/A')}
Journal: {doc.get('journal', 'N/A')}
Year: {doc.get('year', 'N/A')}

Content: {doc['full_content'][:1500]}...
""")
    
    context = "\n".join(context_parts)
    
    # Create prompt
    prompt = f"""You are a research assistant specializing in e-cigarette and vaping research. Based on the provided research documents, answer the user's question accurately and comprehensively.

Context from research documents:
{context}

User Question: {question}

Instructions:
1. Provide a detailed, evidence-based answer using the information from the provided documents
2. Cite specific findings, statistics, or conclusions from the research
3. If the documents contain conflicting information, acknowledge this
4. Be precise about what the research shows vs. what it doesn't address
5. Format your response clearly with key points highlighted
6. If the question cannot be fully answered from the provided documents, state this clearly

Answer:"""

    client = AsyncOpenAI(api_key=api_key)
    
    try:
        response = await client.chat.completions.create(
            model="gpt-4.1",
            messages=[
                {"role": "system", "content": "You are a helpful research assistant specializing in e-cigarette and vaping research. Provide accurate, evidence-based answers based on the provided research documents."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=4096,
            temperature=0.1
        )
        
        return response.choices[0].message.content
    
    except Exception as e:
        return f"Error generating answer: {str(e)}"
    
    finally:
        await client.close()

def process_question(question: str, rag_system: RAGSystem, relevant_documents: List[Dict[str, Any]], api_key: str) -> str:
    """
    Process a question using the RAG system with OpenAI API
    
    Parameters:
    -----------
    question : str
        User question
    rag_system : RAGSystem
        Initialized RAG system
    relevant_documents : List[Dict[str, Any]]
        List of relevant documents
    api_key : str
        OpenAI API key
    
    Returns:
    --------
    str
        Generated answer
    """
    try:
        # Use asyncio to run the async function
        return asyncio.run(generate_answer_async(question, relevant_documents, api_key))
    except RuntimeError as e:
        if "cannot be called from a running event loop" in str(e):
            loop = asyncio.get_event_loop()
            return loop.run_until_complete(generate_answer_async(question, relevant_documents, api_key))
        else:
            raise e


def export_chat_to_docx(chat_history: List[Dict[str, str]]) -> io.BytesIO:
    """
    Export chat history to a Word document (.docx) with properly formatted content and tight spacing
    
    Parameters:
    -----------
    chat_history : List[Dict[str, str]]
        List of chat messages with 'role' and 'content' keys
    
    Returns:
    --------
    io.BytesIO
        Word document as binary stream
    
    Note: This function requires the 'python-docx' package.
    Install it with: pip install python-docx
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml.shared import qn
        from typing import List
        import re
    except ImportError:
        # If docx module is not available, return information about the requirement
        error_msg = ("The 'python-docx' package is required for Word document export.\n"
                    "Please install it with: pip install python-docx")
        raise ImportError(error_msg)
    
    # Create a new Document
    doc = Document()
    
    # Set document-wide spacing defaults for ultra-tight layout
    styles = doc.styles
    
    # Modify Normal style for minimal spacing
    normal_style = styles['Normal']
    normal_paragraph_format = normal_style.paragraph_format
    normal_paragraph_format.space_before = Pt(0)
    normal_paragraph_format.space_after = Pt(2)  # Very small space after paragraphs
    normal_paragraph_format.line_spacing = 1.0   # Single line spacing
    
    # Modify heading styles for ultra-tight spacing
    for level in range(1, 4):
        heading_style = styles[f'Heading {level}']
        heading_format = heading_style.paragraph_format
        heading_format.space_before = Pt(2)  # Minimal space before headings
        heading_format.space_after = Pt(2)   # Minimal space after headings
        heading_format.line_spacing = 1.0
    
    # Modify list styles for tight spacing
    try:
        list_bullet_style = styles['List Bullet']
        list_bullet_format = list_bullet_style.paragraph_format
        list_bullet_format.space_before = Pt(0)
        list_bullet_format.space_after = Pt(1)
        list_bullet_format.line_spacing = 1.0
        
        list_number_style = styles['List Number']
        list_number_format = list_number_style.paragraph_format
        list_number_format.space_before = Pt(0)
        list_number_format.space_after = Pt(1)
        list_number_format.line_spacing = 1.0
    except:
        pass  # Styles might not exist
    
    # Add title
    title = doc.add_heading('E-Cigarette Research Q&A Chat History', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_format = title.paragraph_format
    title_format.space_before = Pt(0)
    title_format.space_after = Pt(6)
    
    # Add subtitle with date
    from datetime import datetime
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    subtitle_run.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.paragraph_format
    subtitle_format.space_before = Pt(0)
    subtitle_format.space_after = Pt(12)
    
    def set_minimal_spacing(paragraph, space_before=0, space_after=2):
        """Set ultra-minimal spacing for paragraphs"""
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(space_before)
        paragraph_format.space_after = Pt(space_after)
        paragraph_format.line_spacing = 1.0
    
    def create_table_from_markdown(table_lines: List[str]):
        """Create a Word table from markdown table lines"""
        if not table_lines:
            return
        
        # Filter out separator lines
        data_lines = []
        for line in table_lines:
            if not re.match(r'^[\|\-\:\s]+$', line):
                data_lines.append(line)
        
        if not data_lines:
            return
        
        # Parse table data
        table_data = []
        for line in data_lines:
            cells = [cell.strip() for cell in line.split('|')]
            if cells and not cells[0]:
                cells = cells[1:]
            if cells and not cells[-1]:
                cells = cells[:-1]
            if cells:
                table_data.append(cells)
        
        if not table_data:
            return
        
        # Determine table dimensions
        max_cols = max(len(row) for row in table_data) if table_data else 0
        if max_cols == 0:
            return
        
        # Ensure all rows have same number of columns
        for row in table_data:
            while len(row) < max_cols:
                row.append('')
        
        # Create table
        table = doc.add_table(rows=len(table_data), cols=max_cols)
        table.style = 'Table Grid'
        
        # Fill table data
        for row_idx, row_data in enumerate(table_data):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell_paragraph = cell.paragraphs[0]
                set_minimal_spacing(cell_paragraph, space_before=0, space_after=0)
                
                # Format first row as header
                if row_idx == 0:
                    run = cell_paragraph.add_run(cell_data)
                    run.bold = True
                    # Add light gray background for header
                    try:
                        from docx.oxml import parse_xml
                        shading_elm = parse_xml(r'<w:shd {} w:fill="E6E6E6"/>'.format(
                            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'))
                        cell._tc.get_or_add_tcPr().append(shading_elm)
                    except:
                        pass
                else:
                    add_formatted_text_to_paragraph(cell_data, cell_paragraph)
    
    def add_formatted_text_to_paragraph(text: str, paragraph_obj):
        """Add text with inline markdown formatting to a paragraph"""
        # Handle bold text (**text**)
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                bold_text = part[2:-2]
                run = paragraph_obj.add_run(bold_text)
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                italic_text = part[1:-1]
                run = paragraph_obj.add_run(italic_text)
                run.italic = True
            else:
                paragraph_obj.add_run(part)
    
    def parse_markdown_to_docx(content: str):
        """Parse markdown content and add formatted text to document with minimal spacing"""
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Skip empty lines entirely - no extra spacing
            if not line:
                i += 1
                continue
            
            # Skip all types of horizontal rules and separators
            if (line.startswith('---') or line.startswith('–') or line.startswith('─') or 
                line == '─' * len(line) or re.match(r'^[-─–]{3,}$', line) or
                line.startswith('\\-\\-\\-')):
                i += 1
                continue
            
            # Handle tables
            if '|' in line and line.count('|') >= 2:
                table_lines = []
                while i < len(lines) and '|' in lines[i]:
                    table_lines.append(lines[i].strip())
                    i += 1
                
                if table_lines:
                    create_table_from_markdown(table_lines)
                continue
            
            # Handle headers with minimal spacing
            if line.startswith('###'):
                header_text = line.replace('###', '').strip()
                header = doc.add_heading(header_text, level=3)
                set_minimal_spacing(header, space_before=2, space_after=2)
                i += 1
                continue
            elif line.startswith('##'):
                header_text = line.replace('##', '').strip()
                header = doc.add_heading(header_text, level=2)
                set_minimal_spacing(header, space_before=2, space_after=2)
                i += 1
                continue
            elif line.startswith('#'):
                header_text = line.replace('#', '').strip()
                header = doc.add_heading(header_text, level=1)
                set_minimal_spacing(header, space_before=2, space_after=2)
                i += 1
                continue
            
            # Handle bullet points
            if line.startswith('- ') or line.startswith('• '):
                bullet_text = line[2:].strip()
                bullet_para = doc.add_paragraph(style='List Bullet')
                set_minimal_spacing(bullet_para, space_before=0, space_after=1)
                add_formatted_text_to_paragraph(bullet_text, bullet_para)
                i += 1
                continue
            
            # Handle numbered lists
            if re.match(r'^\d+\.\s', line):
                numbered_text = re.sub(r'^\d+\.\s', '', line)
                numbered_para = doc.add_paragraph(style='List Number')
                set_minimal_spacing(numbered_para, space_before=0, space_after=1)
                add_formatted_text_to_paragraph(numbered_text, numbered_para)
                i += 1
                continue
            
            # Regular text with inline formatting
            para = doc.add_paragraph()
            set_minimal_spacing(para, space_before=0, space_after=2)
            add_formatted_text_to_paragraph(line, para)
            i += 1
    
    # Process chat messages
    qa_number = 1
    for i, message in enumerate(chat_history):
        if message["role"] == "user":
            # User question with minimal spacing
            user_para = doc.add_paragraph()
            set_minimal_spacing(user_para, space_before=0, space_after=4)
            user_run = user_para.add_run(f"Question {qa_number}: ")
            user_run.bold = True
            user_run.font.size = Pt(12)
            user_run.font.color.rgb = RGBColor(0, 102, 204)  # Blue color
            user_para.add_run(message["content"])
            
        else:  # assistant message
            # Research Bot response
            bot_para = doc.add_paragraph()
            set_minimal_spacing(bot_para, space_before=0, space_after=4)
            bot_run = bot_para.add_run(f"Research Bot Response {qa_number}: ")
            bot_run.bold = True
            bot_run.font.size = Pt(12)
            bot_run.font.color.rgb = RGBColor(34, 139, 34)  # Green color
            
            # Parse and format the markdown content
            parse_markdown_to_docx(message["content"])
            
            # Add sources if available
            if "sources" in message and message["sources"]:
                # Add sources header with minimal spacing
                sources_header = doc.add_paragraph()
                set_minimal_spacing(sources_header, space_before=4, space_after=2)
                sources_run = sources_header.add_run("📚 Sources used for this response:")
                sources_run.bold = True
                sources_run.font.size = Pt(11)
                
                # Add each source
                for j, doc_source in enumerate(message["sources"]):
                    source_para = doc.add_paragraph(style='List Bullet')
                    set_minimal_spacing(source_para, space_before=0, space_after=1)
                    source_title = doc_source.get('title', f'Source {j+1}')
                    source_score = doc_source.get('score', 'N/A')
                    
                    # Add source title in bold
                    title_run = source_para.add_run(source_title)
                    title_run.bold = True
                    
                    # Add similarity score
                    source_para.add_run(f" (Similarity: {source_score})")
                    
                    # Add additional metadata if available
                    metadata_parts = []
                    if doc_source.get('year') and doc_source.get('year') != 'N/A':
                        metadata_parts.append(f"Year: {doc_source.get('year')}")
                    if doc_source.get('authors') and doc_source.get('authors') != 'N/A':
                        authors = doc_source.get('authors')
                        # Format authors properly
                        if isinstance(authors, str) and authors.startswith('[') and authors.endswith(']'):
                            try:
                                import ast
                                authors_list = ast.literal_eval(authors)
                                if isinstance(authors_list, list):
                                    formatted_authors = ', '.join(str(author).strip("'\"") for author in authors_list)
                                else:
                                    formatted_authors = str(authors)
                            except:
                                formatted_authors = authors.strip('[]').replace("'", "").replace('"', '')
                        elif isinstance(authors, list):
                            formatted_authors = ', '.join(str(author) for author in authors)
                        else:
                            formatted_authors = str(authors)
                        metadata_parts.append(f"Authors: {formatted_authors}")
                    
                    if metadata_parts:
                        source_para.add_run(f" | {' | '.join(metadata_parts)}")
            
            qa_number += 1
            
            # Add minimal separator between Q&A pairs (but not after the last one)
            if i < len(chat_history) - 1:
                separator_para = doc.add_paragraph()
                set_minimal_spacing(separator_para, space_before=8, space_after=4)
                separator_run = separator_para.add_run("─" * 50)
                separator_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray color
    
    # Save document to a BytesIO object
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes


# Quick search function for direct usage
def quick_search(query: str, openai_api_key: str, index_path: str = "faiss_index", 
                embedding_model: str = "text-embedding-3-large", top_k: int = 8) -> List[Dict]:
    """
    Quick search function for direct usage.
    
    Args:
        query (str): Search query
        openai_api_key (str): OpenAI API key
        index_path (str): Path to the saved index directory
        embedding_model (str): OpenAI embedding model to use
        top_k (int): Number of top results to return
        
    Returns:
        List[Dict]: List of search results
    """
    # Load index if not already loaded
    index, page_metadata = load_faiss_index(index_path)
    
    if index is None or page_metadata is None:
        return []
    
    # Create temporary RAG system
    rag_system = RAGSystem()
    rag_system.faiss_index = index
    rag_system.page_metadata = page_metadata
    rag_system.is_initialized = True
    
    # Perform search
    results = search_documents_faiss(query, rag_system, openai_api_key, top_k)
    return results